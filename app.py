"""
═══════════════════════════════════════════════════════════════════════════════
 EduRAG — Streamlit UI
═══════════════════════════════════════════════════════════════════════════════
 Three tabs:
   1. Ask — Query your documents with any RAG variant (single or multi-question)
   2. Compare — Side-by-side Vanilla vs Full System
   3. Ablation — Run your questions through selected variants, see metrics
═══════════════════════════════════════════════════════════════════════════════
"""

import streamlit as st
import numpy as np
import pandas as pd
import time
import json
import os
import sys
import io
import re
import logging
import traceback

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from config import RAGConfig, Document, RAGResponse, ABLATION_VARIANTS
from pipeline import RAGPipeline
from utils.metrics import MetricsLogger, QueryLog

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ─────────────────────────────────────────────────────────────────────────────
# Page Config
# ─────────────────────────────────────────────────────────────────────────────
st.set_page_config(page_title="EduRAG", page_icon="📚", layout="wide")

st.markdown("""
<style>
    .metric-box {
        background: #1a1f2e; border: 1px solid #2d3748; border-radius: 10px;
        padding: 16px; margin: 6px 0;
    }
    .metric-box h4 { color: #8b9dc3; font-size: 0.8rem; margin: 0; }
    .metric-box .val { color: #e2e8f0; font-size: 1.5rem; font-weight: 700; }
    .supported { background: #0a2e1a; border-left: 4px solid #22c55e;
        padding: 10px 14px; border-radius: 6px; margin: 6px 0; }
    .unsupported { background: #2e0a0a; border-left: 4px solid #ef4444;
        padding: 10px 14px; border-radius: 6px; margin: 6px 0; }
    .answer-box { background: #1a1f2e; border: 1px solid #2d3748;
        border-radius: 10px; padding: 18px; margin: 10px 0; line-height: 1.7; }
    .question-header { background: #1e293b; border: 1px solid #334155;
        border-radius: 8px; padding: 12px 16px; margin: 16px 0 8px 0;
        font-weight: 600; color: #93c5fd; }
    .eval-badge { display: inline-block; background: #374151; color: #9ca3af;
        font-size: 0.65rem; padding: 2px 6px; border-radius: 4px;
        margin-left: 4px; vertical-align: super; }
</style>
""", unsafe_allow_html=True)


# ─────────────────────────────────────────────────────────────────────────────
# Session State
# ─────────────────────────────────────────────────────────────────────────────
if "documents" not in st.session_state:
    st.session_state.documents = []
if "metrics_logger" not in st.session_state:
    st.session_state.metrics_logger = MetricsLogger("edu_rag")
if "parse_errors" not in st.session_state:
    st.session_state.parse_errors = []
if "ask_export_data" not in st.session_state:
    st.session_state.ask_export_data = []
if "compare_export_data" not in st.session_state:
    st.session_state.compare_export_data = []


# ─────────────────────────────────────────────────────────────────────────────
# Question Parsing Helper
# ─────────────────────────────────────────────────────────────────────────────
def parse_questions(text: str) -> list:
    if not text.strip():
        return []
    lines = text.strip().split("\n")
    questions = []
    for line in lines:
        line = line.strip()
        if not line:
            continue
        cleaned = re.sub(r'^[qQ]?\d+[\.\)\:\-]\s*', '', line).strip()
        if cleaned:
            questions.append(cleaned)
    return questions


# ─────────────────────────────────────────────────────────────────────────────
# Document Parsing
# ─────────────────────────────────────────────────────────────────────────────
def chunk_text(text: str, filename: str, chunk_size: int = 800, overlap: int = 150) -> list:
    docs = []
    text = text.strip()
    if not text:
        return docs
    text = text.replace('\x00', '')
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r' {3,}', ' ', text)

    paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
    current_chunk = ""
    chunk_idx = 0

    for para in paragraphs:
        if len(current_chunk) + len(para) + 1 <= chunk_size:
            current_chunk = f"{current_chunk}\n{para}" if current_chunk else para
        else:
            if current_chunk.strip():
                docs.append(Document(
                    doc_id=f"{filename}_{chunk_idx:04d}",
                    content=current_chunk.strip(),
                    metadata={"source": filename, "chunk": chunk_idx},
                ))
                chunk_idx += 1
            if overlap > 0 and current_chunk:
                overlap_text = current_chunk.strip()[-overlap:]
                current_chunk = f"{overlap_text}\n{para}"
            else:
                current_chunk = para
            while len(current_chunk) > chunk_size:
                split_point = current_chunk[:chunk_size].rfind('. ')
                if split_point == -1 or split_point < chunk_size // 2:
                    split_point = chunk_size
                else:
                    split_point += 1
                docs.append(Document(
                    doc_id=f"{filename}_{chunk_idx:04d}",
                    content=current_chunk[:split_point].strip(),
                    metadata={"source": filename, "chunk": chunk_idx},
                ))
                chunk_idx += 1
                current_chunk = current_chunk[split_point - overlap:].strip()

    if current_chunk.strip():
        docs.append(Document(
            doc_id=f"{filename}_{chunk_idx:04d}",
            content=current_chunk.strip(),
            metadata={"source": filename, "chunk": chunk_idx},
        ))
    return docs


def parse_pdf(file_obj, filename: str) -> list:
    file_obj.seek(0)
    raw_bytes = file_obj.read()
    pages_text = []
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(raw_bytes))
        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                pages_text.append((page_num + 1, text.strip()))
    except (ImportError, Exception):
        pass
    if not pages_text:
        try:
            import fitz
            doc = fitz.open(stream=raw_bytes, filetype="pdf")
            for page_num, page in enumerate(doc):
                text = page.get_text()
                if text and text.strip():
                    pages_text.append((page_num + 1, text.strip()))
            doc.close()
        except (ImportError, Exception):
            pass
    if not pages_text:
        try:
            import pdfplumber
            pdf = pdfplumber.open(io.BytesIO(raw_bytes))
            for page_num, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text and text.strip():
                    pages_text.append((page_num + 1, text.strip()))
            pdf.close()
        except (ImportError, Exception):
            pass
    if not pages_text:
        return []
    full_text = "\n\n".join([text for _, text in pages_text])
    chunks = chunk_text(full_text, filename)
    char_to_page = []
    for page_num, text in pages_text:
        for _ in range(len(text)):
            char_to_page.append(page_num)
    for doc in chunks:
        chunk_start = full_text.find(doc.content[:50])
        if chunk_start >= 0 and chunk_start < len(char_to_page):
            doc.metadata["page"] = char_to_page[min(chunk_start, len(char_to_page) - 1)]
    return chunks


def parse_docx(file_obj, filename: str) -> list:
    file_obj.seek(0)
    try:
        from docx import Document as DocxDocument
        doc_file = DocxDocument(file_obj)
        paragraphs = [p.text.strip() for p in doc_file.paragraphs if p.text.strip()]
        if paragraphs:
            full_text = "\n".join(paragraphs)
            return chunk_text(full_text, filename)
    except (ImportError, Exception):
        pass
    return []


def parse_uploaded_files(uploaded_files) -> tuple:
    docs = []
    errors = []
    for f in uploaded_files:
        name = f.name
        try:
            if name.lower().endswith(".pdf"):
                result = parse_pdf(f, name)
                if result:
                    docs.extend(result)
                else:
                    errors.append(f"❌ **{name}**: No text extracted.")
            elif name.lower().endswith((".docx", ".doc")):
                result = parse_docx(f, name)
                if result:
                    docs.extend(result)
                else:
                    errors.append(f"❌ **{name}**: No text extracted.")
            elif name.lower().endswith(".txt"):
                content = f.read().decode("utf-8", errors="ignore")
                result = chunk_text(content, name)
                docs.extend(result)
            elif name.lower().endswith(".csv"):
                f.seek(0)
                df = pd.read_csv(f)
                if "content" in df.columns:
                    for i, row in df.iterrows():
                        docs.append(Document(
                            doc_id=row.get("doc_id", f"csv_{i:04d}"),
                            content=str(row["content"]),
                            metadata={
                                "source": row.get("source", name),
                                "subject": row.get("subject", ""),
                                "topic": row.get("topic", ""),
                            },
                        ))
                else:
                    errors.append(f"⚠️ **{name}**: CSV must have a 'content' column.")
            elif name.lower().endswith(".json"):
                f.seek(0)
                data = json.loads(f.read())
                if isinstance(data, list):
                    for i, item in enumerate(data):
                        if isinstance(item, dict) and "content" in item:
                            docs.append(Document(
                                doc_id=item.get("doc_id", f"json_{i:04d}"),
                                content=item["content"],
                                metadata=item.get("metadata", {"source": name}),
                            ))
                else:
                    errors.append(f"⚠️ **{name}**: JSON must be a list of objects with 'content' key.")
            else:
                errors.append(f"⚠️ **{name}**: Unsupported file type.")
        except Exception as e:
            errors.append(f"❌ **{name}**: {str(e)}")
    return docs, errors


# ─────────────────────────────────────────────────────────────────────────────
# Pipeline Helper
# ─────────────────────────────────────────────────────────────────────────────
def build_pipeline(config: RAGConfig) -> RAGPipeline:
    pipeline = RAGPipeline(config, st.session_state.metrics_logger)
    pipeline.index_documents(st.session_state.documents)
    return pipeline


# ─────────────────────────────────────────────────────────────────────────────
# Display Helpers
# ─────────────────────────────────────────────────────────────────────────────
def metric_card(label, value, sub=""):
    st.markdown(f"""
    <div class="metric-box">
        <h4>{label}</h4>
        <div class="val">{value}</div>
        <div style="color:#64748b;font-size:0.7rem;">{sub}</div>
    </div>
    """, unsafe_allow_html=True)


def show_claims(claims):
    for i, c in enumerate(claims):
        css = "supported" if c.is_supported else "unsupported"
        icon = "✅" if c.is_supported else "❌"
        score = f"{c.support_score:.3f}" if c.support_score is not None else "N/A"
        st.markdown(f"""
        <div class="{css}">
            <strong>{icon} Claim {i+1}</strong> &nbsp;
            <span style="color:#94a3b8;font-size:0.8rem;">Score: {score}</span>
            <div style="color:#e2e8f0;margin-top:4px;">{c.text}</div>
        </div>
        """, unsafe_allow_html=True)


def show_retrieved_docs(docs):
    for d in docs:
        src = d.document.metadata.get("source", d.document.doc_id)
        page = d.document.metadata.get("page", "")
        page_str = f" (p.{page})" if page else ""
        st.markdown(f"**{src}**{page_str} — similarity: `{d.score:.4f}`")
        st.caption(d.document.content[:300])
        st.markdown("---")


def show_citations(citations, documents):
    if not citations:
        return
    st.markdown("**References:**")
    shown = set()
    for c in citations:
        num = c.get("citation_num", 0)
        doc_id = c.get("doc_id", "")
        if num not in shown:
            source = doc_id
            for doc in documents:
                if doc.document.doc_id == doc_id:
                    source = doc.document.metadata.get("source", doc_id)
                    break
            st.markdown(f"[{num}] {source} (`{doc_id}`)")
            shown.add(num)


def show_all_metrics_row(response):
    """
    Show a unified metrics row for ANY variant.
    Uses get_effective_* so all variants always have scores.
    """
    eff_hall = response.get_effective_hallucination_stats()
    eff_cit = response.get_effective_citation_stats()

    # Determine if scores come from active module or eval pass
    hall_source = "active" if response.hallucination_stats else "eval"
    cit_source = "active" if response.citation_stats else "eval"
    hall_badge = "" if hall_source == "active" else '<span class="eval-badge">eval</span>'
    cit_badge = "" if cit_source == "active" else '<span class="eval-badge">eval</span>'

    cols = st.columns(5)
    with cols[0]:
        metric_card("Retrieval", f"{response.metrics.get('avg_similarity', 0):.4f}",
                     "avg cosine sim")
    with cols[1]:
        metric_card("Total Time", f"{response.module_timings.get('total', 0):.0f}ms",
                     "end-to-end")
    with cols[2]:
        hr = eff_hall.get("hallucination_rate")
        faith = eff_hall.get("faithfulness_score")
        metric_card(
            f"Halluc. Rate",
            f"{hr:.1%}" if hr is not None else "—",
            f"faithfulness: {faith:.3f}" if faith is not None else ""
        )
        if hall_badge:
            st.markdown(f'<div style="text-align:center;margin-top:-8px;">{hall_badge}</div>',
                        unsafe_allow_html=True)
    with cols[3]:
        fs = eff_hall.get("faithfulness_score")
        nc = eff_hall.get("num_claims", 0)
        ns = eff_hall.get("num_supported", 0)
        metric_card(
            "Faithfulness",
            f"{fs:.3f}" if fs is not None else "—",
            f"{ns}/{nc} claims supported"
        )
        if hall_badge:
            st.markdown(f'<div style="text-align:center;margin-top:-8px;">{hall_badge}</div>',
                        unsafe_allow_html=True)
    with cols[4]:
        cc = eff_cit.get("citation_coverage")
        ca = eff_cit.get("avg_confidence", eff_cit.get("citation_accuracy"))
        metric_card(
            "Citation Cov.",
            f"{cc:.1%}" if cc is not None else "—",
            f"confidence: {ca:.3f}" if ca is not None else ""
        )
        if cit_badge:
            st.markdown(f'<div style="text-align:center;margin-top:-8px;">{cit_badge}</div>',
                        unsafe_allow_html=True)


def build_export_record(variant_name, question, response, ground_truth=None):
    """Build a structured dict for export from a single response."""
    eff_hall = response.get_effective_hallucination_stats()
    eff_cit = response.get_effective_citation_stats()

    record = {
        "variant": variant_name,
        "question": question,
        "raw_answer": response.raw_answer,
        "final_answer": response.final_answer,
        "avg_similarity": round(response.metrics.get("avg_similarity", 0), 4),
        "total_time_ms": round(response.module_timings.get("total", 0), 1),
        "answer_length": len(response.final_answer),
        "hallucination_rate": round(eff_hall.get("hallucination_rate", 0), 4),
        "faithfulness_score": round(eff_hall.get("faithfulness_score", 0), 4),
        "num_claims": int(eff_hall.get("num_claims", 0)),
        "num_supported": int(eff_hall.get("num_supported", 0)),
        "num_unsupported": int(eff_hall.get("num_unsupported", 0)),
        "citation_coverage": round(eff_cit.get("citation_coverage", 0), 4),
        "citation_accuracy": round(eff_cit.get("citation_accuracy", eff_cit.get("avg_confidence", 0)), 4),
        "mmr_active": response.mmr_reranked_docs is not None,
        "hallucination_module_active": bool(response.hallucination_stats),
        "citation_module_active": bool(response.citation_stats),
    }

    if ground_truth:
        ml = st.session_state.metrics_logger
        aq = ml.compute_answer_quality(response.final_answer, ground_truth)
        record["ground_truth"] = ground_truth
        record["f1_score"] = round(aq.get("f1_score", 0), 4)
        record["exact_match"] = round(aq.get("exact_match", 0), 4)
        record["rouge_l"] = round(aq.get("rouge_l", 0), 4)

    return record


def build_export_json(records: list) -> str:
    """Build a nicely formatted JSON export grouped by variant."""
    grouped = {}
    for r in records:
        v = r["variant"]
        if v not in grouped:
            grouped[v] = []
        grouped[v].append(r)

    export = {
        "experiment": "EduRAG Results",
        "total_queries": len(records),
        "variants": {}
    }
    for variant_name, variant_records in grouped.items():
        export["variants"][variant_name] = {
            "num_queries": len(variant_records),
            "results": variant_records,
        }
    return json.dumps(export, indent=2, default=str)


def build_export_csv(records: list) -> str:
    df = pd.DataFrame(records)
    return df.to_csv(index=False)


def show_export_buttons(records: list, key_prefix: str):
    """Show CSV and JSON download buttons for a list of export records."""
    if not records:
        return
    st.markdown("### 💾 Export Results")
    col1, col2 = st.columns(2)
    with col1:
        st.download_button(
            "📥 Download CSV", build_export_csv(records),
            f"{key_prefix}_results.csv", "text/csv",
            use_container_width=True, key=f"{key_prefix}_csv"
        )
    with col2:
        st.download_button(
            "📥 Download JSON", build_export_json(records),
            f"{key_prefix}_results.json", "application/json",
            use_container_width=True, key=f"{key_prefix}_json"
        )


def show_single_response(response, variant_name, query_num=None):
    """Display a complete response for one question — always showing ALL metrics."""
    # ── Answer ──
    st.markdown(f'<div class="answer-box">{response.final_answer}</div>',
                unsafe_allow_html=True)
    st.caption(f"🤖 Generated via **{llm_provider}** / `{llm_model}` — Max tokens: {max_tokens}")

    # ── Unified Metrics Row (always shows all scores) ──
    show_all_metrics_row(response)

    # ── Claim-Level Analysis (always available via eval pass) ──
    eff_claims = response.get_effective_claims()
    if eff_claims:
        source_label = "Active Module" if response.claims else "Evaluation Pass"
        with st.expander(f"🔬 Claim-Level Analysis ({source_label})"):
            show_claims(eff_claims)

    # ── Citations (only if active module produced them) ──
    if response.citations:
        with st.expander("📎 Citations"):
            show_citations(response.citations, response.get_active_docs())

    # ── Retrieved Documents ──
    with st.expander("📄 Retrieved Documents"):
        show_retrieved_docs(response.get_active_docs())

    # ── Raw Metrics ──
    with st.expander("📊 Full Metrics"):
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("**Module Timings:**")
            st.json(response.module_timings)
        with col2:
            st.markdown("**All Metrics:**")
            st.json(response.metrics)

        eff_hall = response.get_effective_hallucination_stats()
        if eff_hall:
            st.markdown("**Hallucination Stats (effective):**")
            st.json(eff_hall)

        eff_cit = response.get_effective_citation_stats()
        if eff_cit:
            st.markdown("**Citation Stats (effective):**")
            st.json(eff_cit)


def get_sidebar_config() -> dict:
    return dict(
        top_k=top_k,
        llm_provider=llm_provider,
        llm_model=llm_model,
        llm_temperature=temperature,
        llm_max_tokens=max_tokens,
        hallucination_threshold=hall_threshold,
        mmr_lambda=mmr_lambda,
    )


def apply_sidebar_config(config: RAGConfig):
    config.top_k = top_k
    config.llm_provider = llm_provider
    config.llm_model = llm_model
    config.llm_temperature = temperature
    config.llm_max_tokens = max_tokens
    config.hallucination_threshold = hall_threshold
    config.mmr_lambda = mmr_lambda
    return config


# ─────────────────────────────────────────────────────────────────────────────
# SIDEBAR
# ─────────────────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("## 📚 EduRAG")
    st.markdown("---")

    st.markdown("### 📄 Upload Documents")
    uploaded = st.file_uploader(
        "Upload your educational content",
        type=["txt", "csv", "json", "pdf", "docx", "doc"],
        accept_multiple_files=True,
        help="Supports: PDF, DOCX, TXT, CSV, JSON",
    )

    if uploaded:
        with st.spinner("Parsing documents..."):
            new_docs, parse_errors = parse_uploaded_files(uploaded)
        if new_docs:
            st.session_state.documents = new_docs
            st.session_state.parse_errors = parse_errors
            st.success(f"✅ Loaded **{len(new_docs)}** chunks from {len(uploaded)} file(s)")
        else:
            st.session_state.documents = []
            st.session_state.parse_errors = parse_errors
        for err in parse_errors:
            st.error(err)
    else:
        if st.session_state.documents:
            st.session_state.documents = []

    num_docs = len(st.session_state.documents)
    if num_docs > 0:
        st.info(f"📊 **{num_docs}** document chunks ready")
        with st.expander("Preview documents"):
            for doc in st.session_state.documents[:8]:
                page = doc.metadata.get("page", "")
                page_str = f" (p.{page})" if page else ""
                st.caption(f"**{doc.doc_id}**{page_str}: {doc.content[:120]}...")
            if num_docs > 8:
                st.caption(f"... and {num_docs - 8} more")

    st.markdown("---")

    st.markdown("### ⚙️ LLM Settings")
    llm_provider = st.selectbox("Provider", ["ollama", "openai"])
    llm_model = st.text_input("Model", value="llama3")
    top_k = st.slider("Top-K retrieval", 1, 20, 5)
    temperature = st.slider("Temperature", 0.0, 1.0, 0.3, 0.05)
    max_tokens = st.slider("Max output tokens", 256, 2048, 1024, 128,
                           help="Controls answer length.")

    st.markdown("---")
    st.markdown("### 🎯 Thresholds")
    hall_threshold = st.slider("Hallucination threshold", 0.0, 1.0, 0.5, 0.05)
    mmr_lambda = st.slider("MMR lambda", 0.0, 1.0, 0.7, 0.05,
                           help="Higher = more relevance, Lower = more diversity")


# ─────────────────────────────────────────────────────────────────────────────
# MAIN CONTENT
# ─────────────────────────────────────────────────────────────────────────────
st.markdown("# 📚 EduRAG — Education RAG System")

has_docs = len(st.session_state.documents) > 0

tab_ask, tab_compare, tab_ablation = st.tabs(["🔍 Ask", "⚔️ Compare", "🧪 Ablation Study"])

NO_DOCS_MSG = """
**No documents loaded yet.** Upload files using the sidebar to get started.

Supported: **PDF**, **DOCX**, **TXT**, **CSV**, **JSON**
"""


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# TAB 1: Ask
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
with tab_ask:
    if not has_docs:
        st.warning(NO_DOCS_MSG)
    else:
        st.markdown("Choose a RAG variant and enter your question(s). All metrics are computed for every variant.")

        variant = st.radio(
            "RAG Variant",
            ["Vanilla RAG", "RAG + MMR", "RAG + Hallucination Detection",
             "RAG + Citation", "Full System"],
            horizontal=True,
            key="ask_variant",
        )

        st.markdown("**Enter your question(s)** — one per line, or use numbering like `q1.`, `1.`, `1)`")
        questions_input = st.text_area(
            "Questions",
            placeholder="q1. What is Newton's second law of motion?\nq2. Explain the structure of benzene.",
            height=120,
            key="ask_q",
        )

        ask_clicked = st.button("🔍 Get Answers", type="primary", key="ask_btn",
                                use_container_width=True)

        if ask_clicked and questions_input.strip():
            questions = parse_questions(questions_input)

            if not questions:
                st.warning("No valid questions found. Please enter at least one question.")
            else:
                config = apply_sidebar_config(ABLATION_VARIANTS[variant])
                try:
                    pipeline = build_pipeline(config)
                except Exception as e:
                    st.error(f"Pipeline error: {e}")
                    st.stop()

                if len(questions) > 1:
                    progress = st.progress(0)
                    st.info(f"Processing **{len(questions)}** question(s) using **{variant}**...")

                export_records = []

                for q_idx, q in enumerate(questions):
                    st.markdown(
                        f'<div class="question-header">❓ Question {q_idx + 1}/{len(questions)}: {q}</div>',
                        unsafe_allow_html=True,
                    )

                    with st.spinner(f"Running {variant} via {llm_provider} ({llm_model})..."):
                        try:
                            response = pipeline.query(q)
                        except Exception as e:
                            st.error(f"Error on question {q_idx + 1}: {e}")
                            continue

                    show_single_response(response, variant, query_num=q_idx + 1)
                    export_records.append(build_export_record(variant, q, response))

                    if len(questions) > 1:
                        progress.progress((q_idx + 1) / len(questions))
                    if q_idx < len(questions) - 1:
                        st.markdown("---")

                if len(questions) > 1:
                    progress.progress(1.0)
                    st.success(f"✅ All {len(questions)} questions answered!")

                # ── Export ──
                if export_records:
                    st.markdown("---")
                    show_export_buttons(export_records, "ask")
                    st.session_state.ask_export_data = export_records

        elif ask_clicked:
            st.warning("Please enter at least one question.")


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# TAB 2: Compare
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
with tab_compare:
    if not has_docs:
        st.warning(NO_DOCS_MSG)
    else:
        st.markdown("Same question(s) through **Vanilla RAG** vs **Full System** side-by-side. All metrics computed for both.")

        cmp_input = st.text_area(
            "Your question(s)",
            placeholder="q1. What is photosynthesis?\nq2. Define electromagnetic induction.",
            height=100,
            key="cmp_q",
        )

        ground_truth = st.text_area(
            "Ground truth answer (optional — enables F1, ROUGE-L, Exact Match). For multiple questions, one per line.",
            placeholder="Paste the correct answer(s) here if you have them...",
            height=80, key="cmp_gt",
        )

        cmp_clicked = st.button("⚔️ Compare", type="primary", key="cmp_btn",
                                use_container_width=True)

        if cmp_clicked and cmp_input.strip():
            cmp_questions = parse_questions(cmp_input)
            gt_lines = [g.strip() for g in ground_truth.strip().split("\n") if g.strip()] if ground_truth.strip() else []

            if not cmp_questions:
                st.warning("No valid questions found.")
            else:
                config_vanilla = RAGConfig(
                    use_mmr=False, use_citation=False, use_hallucination_detection=False,
                    **get_sidebar_config(),
                )
                config_full = RAGConfig(
                    use_mmr=True, use_citation=True, use_hallucination_detection=True,
                    **get_sidebar_config(),
                )

                export_records = []

                for q_idx, cmp_query in enumerate(cmp_questions):
                    gt = gt_lines[q_idx] if q_idx < len(gt_lines) else None

                    if len(cmp_questions) > 1:
                        st.markdown(
                            f'<div class="question-header">❓ Question {q_idx + 1}/{len(cmp_questions)}: {cmp_query}</div>',
                            unsafe_allow_html=True,
                        )

                    col_left, col_right = st.columns(2)

                    with col_left:
                        st.markdown("### 📦 Vanilla RAG")
                        with st.spinner("Running..."):
                            pipe_v = build_pipeline(config_vanilla)
                            resp_v = pipe_v.query(cmp_query, ground_truth=gt)
                        st.markdown(f'<div class="answer-box">{resp_v.final_answer}</div>',
                                    unsafe_allow_html=True)
                        st.caption(f"🤖 Generated via **{llm_provider}** / `{llm_model}`")
                        show_all_metrics_row(resp_v)

                        eff_claims_v = resp_v.get_effective_claims()
                        if eff_claims_v:
                            with st.expander("🔬 Claims (eval)"):
                                show_claims(eff_claims_v)

                    with col_right:
                        st.markdown("### 🛡️ Full System")
                        st.caption("MMR + Hallucination Detection + Citation")
                        with st.spinner("Running..."):
                            pipe_f = build_pipeline(config_full)
                            resp_f = pipe_f.query(cmp_query, ground_truth=gt)
                        st.markdown(f'<div class="answer-box">{resp_f.final_answer}</div>',
                                    unsafe_allow_html=True)
                        st.caption(f"🤖 Generated via **{llm_provider}** / `{llm_model}`")
                        show_all_metrics_row(resp_f)

                        if resp_f.claims:
                            with st.expander("🔬 Claims (active)"):
                                show_claims(resp_f.claims)
                        if resp_f.citations:
                            with st.expander("📎 Citations"):
                                show_citations(resp_f.citations, resp_f.get_active_docs())

                    # ── Build export records ──
                    export_records.append(build_export_record("Vanilla RAG", cmp_query, resp_v, gt))
                    export_records.append(build_export_record("Full System", cmp_query, resp_f, gt))

                    # ── Answer Quality Comparison (if ground truth) ──
                    if gt:
                        st.markdown("---")
                        st.markdown("### 📊 Answer Quality Comparison")
                        ml = st.session_state.metrics_logger
                        aq_v = ml.compute_answer_quality(resp_v.final_answer, gt)
                        aq_f = ml.compute_answer_quality(resp_f.final_answer, gt)

                        eff_hall_v = resp_v.get_effective_hallucination_stats()
                        eff_hall_f = resp_f.get_effective_hallucination_stats()
                        eff_cit_v = resp_v.get_effective_citation_stats()
                        eff_cit_f = resp_f.get_effective_citation_stats()

                        comp_df = pd.DataFrame({
                            "Metric": [
                                "F1 Score", "Exact Match", "ROUGE-L",
                                "Hallucination Rate", "Faithfulness Score",
                                "Citation Coverage", "Citation Accuracy",
                            ],
                            "Vanilla RAG": [
                                f"{aq_v.get('f1_score',0):.4f}",
                                f"{aq_v.get('exact_match',0):.0f}",
                                f"{aq_v.get('rouge_l',0):.4f}",
                                f"{eff_hall_v.get('hallucination_rate',0):.4f}",
                                f"{eff_hall_v.get('faithfulness_score',0):.4f}",
                                f"{eff_cit_v.get('citation_coverage',0):.4f}",
                                f"{eff_cit_v.get('citation_accuracy', eff_cit_v.get('avg_confidence',0)):.4f}",
                            ],
                            "Full System": [
                                f"{aq_f.get('f1_score',0):.4f}",
                                f"{aq_f.get('exact_match',0):.0f}",
                                f"{aq_f.get('rouge_l',0):.4f}",
                                f"{eff_hall_f.get('hallucination_rate',0):.4f}",
                                f"{eff_hall_f.get('faithfulness_score',0):.4f}",
                                f"{eff_cit_f.get('citation_coverage',0):.4f}",
                                f"{eff_cit_f.get('citation_accuracy', eff_cit_f.get('avg_confidence',0)):.4f}",
                            ],
                        })
                        st.dataframe(comp_df, use_container_width=True, hide_index=True)
                    else:
                        # Even without ground truth, show halluc/citation comparison
                        st.markdown("---")
                        st.markdown("### 📊 Metrics Comparison")
                        eff_hall_v = resp_v.get_effective_hallucination_stats()
                        eff_hall_f = resp_f.get_effective_hallucination_stats()
                        eff_cit_v = resp_v.get_effective_citation_stats()
                        eff_cit_f = resp_f.get_effective_citation_stats()

                        comp_df = pd.DataFrame({
                            "Metric": [
                                "Avg Similarity",
                                "Hallucination Rate", "Faithfulness Score",
                                "Num Claims", "Num Supported",
                                "Citation Coverage", "Citation Accuracy",
                                "Total Time (ms)",
                            ],
                            "Vanilla RAG": [
                                f"{resp_v.metrics.get('avg_similarity',0):.4f}",
                                f"{eff_hall_v.get('hallucination_rate',0):.4f}",
                                f"{eff_hall_v.get('faithfulness_score',0):.4f}",
                                f"{eff_hall_v.get('num_claims',0):.0f}",
                                f"{eff_hall_v.get('num_supported',0):.0f}",
                                f"{eff_cit_v.get('citation_coverage',0):.4f}",
                                f"{eff_cit_v.get('citation_accuracy', eff_cit_v.get('avg_confidence',0)):.4f}",
                                f"{resp_v.module_timings.get('total',0):.0f}",
                            ],
                            "Full System": [
                                f"{resp_f.metrics.get('avg_similarity',0):.4f}",
                                f"{eff_hall_f.get('hallucination_rate',0):.4f}",
                                f"{eff_hall_f.get('faithfulness_score',0):.4f}",
                                f"{eff_hall_f.get('num_claims',0):.0f}",
                                f"{eff_hall_f.get('num_supported',0):.0f}",
                                f"{eff_cit_f.get('citation_coverage',0):.4f}",
                                f"{eff_cit_f.get('citation_accuracy', eff_cit_f.get('avg_confidence',0)):.4f}",
                                f"{resp_f.module_timings.get('total',0):.0f}",
                            ],
                        })
                        st.dataframe(comp_df, use_container_width=True, hide_index=True)

                    if q_idx < len(cmp_questions) - 1:
                        st.markdown("---")
                        st.markdown("")

                # ── Export ──
                if export_records:
                    st.markdown("---")
                    show_export_buttons(export_records, "compare")
                    st.session_state.compare_export_data = export_records

        elif cmp_clicked:
            st.warning("Please enter at least one question.")


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# TAB 3: Ablation Study
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
with tab_ablation:
    if not has_docs:
        st.warning(NO_DOCS_MSG)
    else:
        st.markdown("Enter your questions, pick variants, and run the experiment. **All metrics are computed for every variant.**")

        st.markdown("### Step 1: Enter your questions (one per line)")
        questions_text = st.text_area(
            "Questions",
            placeholder="q1. What is Newton's second law of motion?\nq2. Explain Le Chatelier's principle.",
            height=150,
            key="abl_questions",
        )
        questions = parse_questions(questions_text)
        if questions:
            st.success(f"📝 {len(questions)} question(s) parsed")

        st.markdown("### Step 2: Select variants to compare")
        selected_variants = st.multiselect(
            "Variants",
            list(ABLATION_VARIANTS.keys()),
            default=["Vanilla RAG", "Full System"],
            key="abl_variants",
        )

        st.markdown("### Step 3: Run")
        num_runs = st.slider("Runs per question", 1, 5, 1, key="abl_runs")

        ready = len(questions) > 0 and len(selected_variants) > 0

        if not ready:
            if not questions:
                st.info("☝️ Enter at least one question above.")
            if not selected_variants:
                st.info("☝️ Select at least one variant.")

        if st.button("🚀 Run Ablation Study", type="primary", disabled=not ready,
                     use_container_width=True):

            st.session_state.metrics_logger = MetricsLogger("ablation")

            total_steps = len(selected_variants) * len(questions) * num_runs
            progress = st.progress(0)
            status = st.empty()
            step = 0
            results = []
            export_records = []

            for variant_name in selected_variants:
                config = apply_sidebar_config(ABLATION_VARIANTS[variant_name])
                pipeline = build_pipeline(config)

                for q in questions:
                    for run in range(num_runs):
                        step += 1
                        progress.progress(step / total_steps)
                        status.text(f"{variant_name} | {q[:50]}... | Run {run+1}/{num_runs}")

                        try:
                            resp = pipeline.query(q)
                        except Exception as e:
                            logger.error(f"Query failed: {e}")
                            continue

                        eff_hall = resp.get_effective_hallucination_stats()
                        eff_cit = resp.get_effective_citation_stats()

                        row = {
                            "Variant": variant_name,
                            "Question": q,
                            "Run": run + 1,
                            "Avg Similarity": round(resp.metrics.get("avg_similarity", 0), 4),
                            "Total Time (ms)": round(resp.module_timings.get("total", 0), 1),
                            "Answer Length": len(resp.final_answer),
                            # Always present via effective getters
                            "Halluc. Rate": round(eff_hall.get("hallucination_rate", 0), 4),
                            "Faithfulness": round(eff_hall.get("faithfulness_score", 0), 4),
                            "Claims": int(eff_hall.get("num_claims", 0)),
                            "Supported": int(eff_hall.get("num_supported", 0)),
                            "Cit. Coverage": round(eff_cit.get("citation_coverage", 0), 4),
                            "Cit. Confidence": round(eff_cit.get("avg_confidence", eff_cit.get("citation_accuracy", 0)), 4),
                            "MMR Active": resp.mmr_reranked_docs is not None,
                            "Hall. Module Active": bool(resp.hallucination_stats),
                            "Cit. Module Active": bool(resp.citation_stats),
                        }
                        results.append(row)
                        export_records.append(build_export_record(variant_name, q, resp))

            progress.progress(1.0)
            status.text("✅ Done!")

            if not results:
                st.error("No results generated. Check your LLM settings.")
            else:
                st.markdown("---")
                df = pd.DataFrame(results)

                # ── Summary ──
                st.markdown("### 📊 Summary by Variant")
                num_cols = ["Avg Similarity", "Total Time (ms)", "Answer Length",
                            "Halluc. Rate", "Faithfulness",
                            "Cit. Coverage", "Cit. Confidence"]
                existing = [c for c in num_cols if c in df.columns]
                summary = df.groupby("Variant")[existing].agg(["mean", "std"]).round(4)
                st.dataframe(summary, use_container_width=True)

                # ── Charts ──
                st.markdown("### Hallucination Rate by Variant")
                st.bar_chart(df.groupby("Variant")["Halluc. Rate"].mean())

                st.markdown("### Faithfulness Score by Variant")
                st.bar_chart(df.groupby("Variant")["Faithfulness"].mean())

                cit_data = df[df["Cit. Coverage"] > 0]
                if not cit_data.empty:
                    st.markdown("### Citation Coverage by Variant")
                    st.bar_chart(df.groupby("Variant")["Cit. Coverage"].mean())

                # ── Full Table ──
                st.markdown("### Full Results")
                st.dataframe(df, use_container_width=True, hide_index=True)

                # ── Export ──
                st.markdown("---")
                col1, col2 = st.columns(2)
                with col1:
                    st.download_button("📥 CSV (table)", df.to_csv(index=False),
                                       "ablation_results.csv", "text/csv",
                                       use_container_width=True, key="abl_csv")
                with col2:
                    st.download_button("📥 JSON (detailed)",
                                       build_export_json(export_records),
                                       "ablation_results.json", "application/json",
                                       use_container_width=True, key="abl_json")